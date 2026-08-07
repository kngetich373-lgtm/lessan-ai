"""Validation tests for the Professional Document Intelligence System.

Run with:  python3 -m unittest tests.test_documents -v

Covers:
1. Document type taxonomy and natural-language kind resolution.
2. Output-format resolution (explicit + natural language + default).
3. Template selection (explicit -> kind default -> keyword -> generic).
4. Publishing formatting: heading numbering, appendices, captions, TOC,
   and global-reference materialisation.
5. AI content generation (strict-JSON parse) and skeleton fallback.
6. All six export formats (DOCX, PDF, Markdown, HTML, RTF, TXT).
7. The DocumentGenerator facade end-to-end pipeline with memory + events.
8. DI container registration (idempotent) and unregistration.
9. Workflow Engine integration (4-step resolve/draft/format/export).
10. DocumentAgent capabilities.
"""

import json
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any, Dict, List

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from documents import DocumentGenerator, DocumentProject, DocumentRequest, DocumentSection
from documents.agent import DocumentAgent
from documents.content import ContentGenerator
from documents.di import register_document_system, unregister_document_system
from documents.exporters import ExportManager
from documents.formatter import DocumentFormatter
from documents.models import (
    DocumentResult,
    Figure,
    OutputFormat,
    Paragraph,
    Table,
)
from documents.resolver import DocumentResolver
from documents.template_manager import DocumentTemplateManager
from documents.types import DocumentTypeRegistry
from documents.workflow import build_document_workflow, register_document_workflow


# --------------------------------------------------------------------------- #
# Fakes
# --------------------------------------------------------------------------- #
class FakeRouter:
    """Implements the orchestrator ModelRouter interface."""

    def __init__(self, payload: str) -> None:
        self.payload = payload
        self.calls = 0

    def complete(self, prompt: str, *, system=None, max_tokens=512,
                 temperature=0.7, model=None) -> str:
        self.calls += 1
        return self.payload

    def is_available(self) -> bool:
        return True


class FakeMemory:
    """Minimal MemoryStore adapter (load/save)."""

    def __init__(self) -> None:
        self._data: Dict[str, Any] = {}

    def load(self) -> Dict[str, Any]:
        return dict(self._data)

    def save(self, memory_update: Dict[str, Any]) -> Dict[str, Any]:
        self._data.update(memory_update)
        return dict(self._data)


class FakeEventBus:
    def __init__(self) -> None:
        self.events: List[str] = []

    def emit(self, event: str, payload: Dict[str, Any]) -> None:
        self.events.append(event)


def build_sample_project() -> DocumentProject:
    return DocumentProject(
        kind="software_requirements",
        title="Acme Messaging Platform — SRS",
        subtitle="Software Requirements Specification",
        author="Lessan AI",
        sections=[
            DocumentSection(heading="Introduction", paragraphs=[Paragraph("Intro text.")]),
            DocumentSection(heading="Functional Requirements", bullets=["Send messages", "Encryption"]),
            DocumentSection(
                heading="Non-Functional Requirements",
                tables=[Table(rows=[["Metric", "Target"], ["Latency", "<120ms"]], caption="Performance targets")],
            ),
            DocumentSection(heading="Configuration", level=2, paragraphs=[Paragraph("Detail.")]),
            DocumentSection(
                heading="Appendix A: Glossary",
                appendix=True,
                paragraphs=[Paragraph("Glossary.")],
            ),
        ],
        references=["ISO/IEC 29148:2018", "RFC 8446 — TLS 1.3"],
    )


def make_generator(output_dir: Path, **kwargs) -> DocumentGenerator:
    return DocumentGenerator(
        exporters=ExportManager(output_dir=output_dir),
        **kwargs,
    )

# --------------------------------------------------------------------------- #
# 1. Document type taxonomy & kind resolution
# --------------------------------------------------------------------------- #
class TestTypeRegistryResolution(unittest.TestCase):
    def setUp(self) -> None:
        self.registry = DocumentTypeRegistry()

    def test_all_expected_types_registered(self) -> None:
        ids = {t.type_id for t in self.registry.all()}
        expected = {
            "research_proposal", "thesis", "resume", "cover_letter",
            "business_plan", "technical_documentation", "software_requirements",
            "software_design", "project_report", "meeting_minutes", "user_manual",
            "api_documentation", "presentation", "letter", "invoice", "quotation",
            "report",
        }
        self.assertTrue(expected.issubset(ids), msg=f"missing: {expected - ids}")

    def test_explicit_id(self) -> None:
        self.assertEqual(self.registry.resolve("invoice").type_id, "invoice")

    def test_alias_resolution(self) -> None:
        self.assertEqual(self.registry.resolve("cv").type_id, "resume")
        self.assertEqual(self.registry.resolve("srs").type_id, "software_requirements")

    def test_natural_language_keyword_scoring(self) -> None:
        cases = {
            "Need a CV and a resume draft": "resume",
            "invoice please": "invoice",
            "make a proposal for our new API": "research_proposal",
            "project report status update": "project_report",
            "api docs for the payment gateway": "api_documentation",
            "meeting minutes from today": "meeting_minutes",
            "business plan for a cafe": "business_plan",
            "write my PhD thesis": "thesis",
            "cover letter for the job": "cover_letter",
            "user manual for the coffee machine": "user_manual",
        }
        for text, expected in cases.items():
            with self.subTest(text=text):
                self.assertEqual(self.registry.resolve(text).type_id, expected)

    def test_specific_beats_generic(self) -> None:
        # "project report" must win over the generic "report".
        self.assertEqual(self.registry.resolve("send me the project report").type_id, "project_report")

    def test_empty_falls_back_to_report(self) -> None:
        self.assertEqual(self.registry.resolve(None).type_id, "report")
        self.assertEqual(self.registry.resolve("").type_id, "report")


# --------------------------------------------------------------------------- #
# 2. Format resolution
# --------------------------------------------------------------------------- #
class TestFormatResolution(unittest.TestCase):
    def setUp(self) -> None:
        self.resolver = DocumentResolver(DocumentTypeRegistry())

    def test_explicit_format(self) -> None:
        self.assertEqual(
            self.resolver.resolve_format(explicit="pdf"),
            [OutputFormat.PDF],
        )

    def test_explicit_format_list(self) -> None:
        self.assertEqual(
            self.resolver.resolve_format(explicit=["docx", "pdf", "md"]),
            [OutputFormat.DOCX, OutputFormat.PDF, OutputFormat.MD],
        )

    def test_natural_language_formats(self) -> None:
        self.assertEqual(
            self.resolver.resolve_format(text="please export as pdf and word doc"),
            [OutputFormat.PDF, OutputFormat.DOCX],
        )

    def test_default_is_docx(self) -> None:
        self.assertEqual(
            self.resolver.resolve_format(text="just write it"),
            [OutputFormat.DOCX],
        )


# --------------------------------------------------------------------------- #
# 3. Template selection
# --------------------------------------------------------------------------- #
class TestTemplateSelection(unittest.TestCase):
    def setUp(self) -> None:
        self.templates = DocumentTemplateManager()

    def test_explicit_template(self) -> None:
        selected = self.templates.select(explicit="academic", text="")
        self.assertEqual(selected.template_id, "academic")

    def test_kind_default_template(self) -> None:
        selected = self.templates.select(
            explicit=None, kind_default="academic", text="generic text"
        )
        self.assertEqual(selected.template_id, "academic")

    def test_keyword_scoring_falls_through(self) -> None:
        selected = self.templates.select(explicit=None, kind_default=None, text="academic paper")
        self.assertEqual(selected.template_id, "academic")

    def test_generic_fallback(self) -> None:
        selected = self.templates.select(explicit=None, kind_default=None, text="")
        self.assertEqual(selected.template_id, "generic")

    def test_unknown_explicit_falls_back(self) -> None:
        selected = self.templates.select(explicit="does_not_exist", text="")
        self.assertIsNotNone(selected)


# --------------------------------------------------------------------------- #
# 4. Publishing formatting
# --------------------------------------------------------------------------- #
class TestFormatter(unittest.TestCase):
    def setUp(self) -> None:
        self.formatter = DocumentFormatter()

    def test_heading_numbering_and_appendix(self) -> None:
        project = build_sample_project()
        formatted = self.formatter.format(project, None)
        numbers = [s.meta["number"] for s in formatted.project.sections]
        self.assertEqual(numbers, ["1", "2", "3", "3.1", "Appendix A", "4"])

    def test_table_and_figure_captions(self) -> None:
        project = DocumentProject(
            kind="report",
            title="T",
            sections=[
                DocumentSection(heading="H1", tables=[
                    Table(rows=[["a", "b"]], caption="First table"),
                ]),
                DocumentSection(heading="H2", figures=[
                    Figure(caption="First figure", path="x.png"),
                ]),
                DocumentSection(heading="H3", tables=[
                    Table(rows=[["c", "d"]], caption="Second table"),
                ]),
            ],
        )
        formatted = self.formatter.format(project, None)
        captions = [t.caption for s in formatted.project.sections for t in s.tables]
        self.assertEqual(captions, ["Table 1: First table", "Table 2: Second table"])
        fig_captions = [f.caption for s in formatted.project.sections for f in s.figures]
        self.assertEqual(fig_captions, ["Figure 1: First figure"])

    def test_toc_entries(self) -> None:
        project = build_sample_project()
        formatted = self.formatter.format(project, None)
        labels = [entry.label for entry in formatted.toc_entries]
        self.assertIn("1  Introduction", labels)
        self.assertIn("3.1  Configuration", labels)
        # The materialised References section is a real section, so it appears
        # in the TOC; appendices are excluded.
        self.assertTrue(any("References" in label for label in labels))
        self.assertFalse(any("Appendix" in label for label in labels))

    def test_global_references_materialised_once(self) -> None:
        project = build_sample_project()
        formatted = self.formatter.format(project, None)
        ref_sections = [
            s for s in formatted.project.sections
            if s.heading.strip().lower().startswith("reference")
        ]
        self.assertEqual(len(ref_sections), 1)
        self.assertEqual(
            ref_sections[0].references,
            ["ISO/IEC 29148:2018", "RFC 8446 — TLS 1.3"],
        )
        self.assertTrue(formatted.has_references)


# --------------------------------------------------------------------------- #
# 5. Content generation (AI + fallback)
# --------------------------------------------------------------------------- #
class TestContentGenerator(unittest.TestCase):
    def setUp(self) -> None:
        self.doc_type = DocumentTypeRegistry().get("invoice")

    def test_ai_json_path(self) -> None:
        payload = {
            "title": "Neural Compression Survey",
            "subtitle": "An overview of learned codecs",
            "sections": [
                {"heading": "Background", "paragraphs": ["Para one.", "Para two."]},
                {"heading": "Methods", "bullets": ["Hyperprior"], "numbered": ["Step A"]},
                {
                    "heading": "Results",
                    "tables": [{"caption": "BD-rate", "rows": [["Codec", "BD-rate"], ["A", "-5%"]]}],
                    "code": {"language": "python", "text": "x = 1"},
                },
            ],
            "references": ["Ballé et al. 2018"],
            "metadata": {"invoice_number": "INV-99"},
        }
        generator = ContentGenerator(model_router=FakeRouter(json.dumps(payload)))
        request = DocumentRequest(document_type="invoice", topic="Compression", formats=[OutputFormat.DOCX])
        project = generator.generate(request, self.doc_type, {})
        self.assertEqual(project.metadata.get("generated_from"), "ai")
        self.assertEqual(project.metadata.get("invoice_number"), "INV-99")
        self.assertEqual(len(project.sections), 3)
        self.assertEqual(sum(len(s.code_blocks) for s in project.sections), 1)
        self.assertEqual(sum(len(s.tables) for s in project.sections), 1)
        self.assertIn("Ballé et al. 2018", project.references)

    def test_malformed_json_then_repair(self) -> None:
        router = FakeRouter("this is not json")
        generator = ContentGenerator(model_router=router)
        request = DocumentRequest(document_type="invoice", topic="Invoice", formats=[OutputFormat.DOCX])
        project = generator.generate(request, self.doc_type, {})
        # One retry with the repair prompt, then skeleton fallback.
        self.assertEqual(router.calls, 2)
        self.assertEqual(project.metadata.get("generated_from"), "skeleton")

    def test_no_router_skeleton_fallback(self) -> None:
        generator = ContentGenerator(model_router=None)
        request = DocumentRequest(
            document_type="invoice",
            topic="Invoice for services",
            content="Item A: 50\nItem B: 100",
            formats=[OutputFormat.DOCX],
        )
        project = generator.generate(request, self.doc_type, {})
        self.assertEqual(project.metadata.get("generated_from"), "skeleton")
        self.assertGreaterEqual(len(project.sections), 1)
        self.assertIn("Invoice for services", project.title)
        # User content lands in the first section.
        joined = " ".join(p.text for p in project.sections[0].paragraphs)
        self.assertIn("Item A: 50", joined)


# --------------------------------------------------------------------------- #
# 6. Export formats
# --------------------------------------------------------------------------- #
class TestExports(unittest.TestCase):
    ALL_FORMATS = [
        OutputFormat.DOCX, OutputFormat.PDF, OutputFormat.MD,
        OutputFormat.HTML, OutputFormat.RTF, OutputFormat.TXT,
    ]

    def test_all_six_formats_exported(self) -> None:
        with TemporaryDirectory() as tmp:
            formatter = DocumentFormatter()
            exporter = ExportManager(output_dir=Path(tmp))
            formatted = formatter.format(build_sample_project(), None)
            for fmt in self.ALL_FORMATS:
                with self.subTest(fmt=fmt.value):
                    path = exporter.export(formatted, fmt, output_name="all_formats")
                    self.assertTrue(Path(path).exists(), f"{fmt.value} not written")
                    self.assertGreater(Path(path).stat().st_size, 0)

    def test_docx_contains_numbering_and_references(self) -> None:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            self.skipTest("python-docx not installed")
        with TemporaryDirectory() as tmp:
            formatter = DocumentFormatter()
            exporter = ExportManager(output_dir=Path(tmp))
            formatted = formatter.format(build_sample_project(), None)
            path = exporter.export(formatted, OutputFormat.DOCX, output_name="srs_docx")
            doc = DocxDocument(path)
            headings = [p.text for p in doc.paragraphs
                        if p.style.name.startswith("Heading") and p.text.strip()]
            self.assertIn("1  Introduction", headings)
            self.assertIn("3.1  Configuration", headings)
            self.assertIn("Appendix A  Appendix A: Glossary", headings)
            refs = [p.text for p in doc.paragraphs if "29148" in p.text]
            self.assertTrue(refs, "global references missing from DOCX")

    def test_pdf_contains_all_sections(self) -> None:
        try:
            import fitz  # type: ignore
        except ImportError:
            self.skipTest("PyMuPDF not installed")
        with TemporaryDirectory() as tmp:
            formatter = DocumentFormatter()
            exporter = ExportManager(output_dir=Path(tmp))
            formatted = formatter.format(build_sample_project(), None)
            path = exporter.export(formatted, OutputFormat.PDF, output_name="srs_pdf")
            text = " ".join(page.get_text() for page in fitz.open(path))
            self.assertIn("Introduction", text)
            self.assertIn("Appendix A", text)
            self.assertIn("29148", text)

    def test_markdown_renders_references_once(self) -> None:
        with TemporaryDirectory() as tmp:
            formatter = DocumentFormatter()
            exporter = ExportManager(output_dir=Path(tmp))
            formatted = formatter.format(build_sample_project(), None)
            path = exporter.export(formatted, OutputFormat.MD, output_name="srs_md")
            text = Path(path).read_text(encoding="utf-8")
            self.assertEqual(text.count("29148"), 1)
            self.assertEqual(text.count("8446"), 1)


# --------------------------------------------------------------------------- #
# 7. DocumentGenerator facade
# --------------------------------------------------------------------------- #
class TestDocumentGenerator(unittest.TestCase):
    def test_full_pipeline_skeleton(self) -> None:
        with TemporaryDirectory() as tmp:
            memory = FakeMemory()
            bus = FakeEventBus()
            gen = make_generator(
                Path(tmp),
                memory_store=memory,
                event_bus=bus,
            )
            request = DocumentRequest(
                document_type="invoice",
                topic="Q3 Services Invoice",
                content="Design services: 500",
                formats=[OutputFormat.DOCX, OutputFormat.TXT],
                output_name="q3_invoice",
            )
            result = gen.generate(request)
            self.assertIsInstance(result, DocumentResult)
            self.assertEqual(result.kind, "invoice")
            self.assertEqual(result.generated_from, "skeleton")
            self.assertEqual(set(result.paths), {OutputFormat.DOCX, OutputFormat.TXT})
            for path in result.paths.values():
                self.assertTrue(Path(path).exists())
            # Memory record written.
            self.assertIn("documents.last", memory.load())
            # Events emitted through the pipeline.
            for expected in (
                "documents.started", "documents.kind_resolved",
                "documents.template_selected", "documents.content_ready",
                "documents.formatted", "documents.exported", "documents.generated",
            ):
                self.assertIn(expected, bus.events)

    def test_natural_language_request_no_explicit_type(self) -> None:
        with TemporaryDirectory() as tmp:
            gen = make_generator(Path(tmp))
            result = gen.generate(
                DocumentRequest(
                    topic="Please draft my resume for a senior role",
                    formats=[OutputFormat.DOCX],
                )
            )
            self.assertEqual(result.kind, "resume")

    def test_introspection_lists(self) -> None:
        gen = DocumentGenerator()
        self.assertIn("invoice", gen.list_document_types())
        self.assertIn("academic", gen.list_templates())

    def test_generate_from_project(self) -> None:
        with TemporaryDirectory() as tmp:
            gen = make_generator(Path(tmp))
            result = gen.generate_from_project(
                DocumentRequest(
                    document_type="meeting_minutes",
                    formats=[OutputFormat.PDF],
                    output_name="standup_minutes",
                ),
                build_sample_project(),
            )
            # The project carries its own kind, which is authoritative for
            # generate_from_project.
            self.assertEqual(result.kind, "software_requirements")
            self.assertIn(OutputFormat.PDF, result.paths)


# --------------------------------------------------------------------------- #
# 8. DI container integration
# --------------------------------------------------------------------------- #
class TestDI(unittest.TestCase):
    def test_register_is_idempotent(self) -> None:
        with TemporaryDirectory() as tmp:
            from core.di.container import Container

            container = Container()
            register_document_system(container, config=None, event_bus=None)
            first = container.resolve(DocumentGenerator)
            register_document_system(container)  # second call must be a no-op
            second = container.resolve(DocumentGenerator)
            self.assertIs(first, second)

    def test_unregister(self) -> None:
        from core.di.container import Container

        container = Container()
        register_document_system(container)
        self.assertTrue(container.has(DocumentGenerator))
        unregister_document_system(container)
        self.assertFalse(container.has(DocumentGenerator))


# --------------------------------------------------------------------------- #
# 9. Workflow Engine integration
# --------------------------------------------------------------------------- #
class TestWorkflowIntegration(unittest.TestCase):
    def test_build_document_workflow(self) -> None:
        workflow = build_document_workflow()
        self.assertEqual(workflow.name, "document_generation")
        self.assertEqual(len(workflow.steps), 4)
        actions = [s.action for s in workflow.steps]
        self.assertEqual(
            actions,
            ["documents.resolve", "documents.draft", "documents.format", "documents.export"],
        )
        self.assertEqual(workflow.steps[1].depends_on, ["resolve"])
        self.assertEqual(workflow.steps[2].depends_on, ["draft"])
        self.assertEqual(workflow.steps[3].depends_on, ["format"])

    def test_workflow_class_and_engine_registration(self) -> None:
        from core.workflow.engine import WorkflowEngine

        engine = WorkflowEngine()
        register_document_workflow(engine)
        self.assertIn("document_generation", engine.registry.list_workflows())
        wf = engine.registry.get("document_generation")()
        self.assertEqual(wf.name, "document_generation")
        self.assertEqual(len(wf.steps), 4)

    def test_request_serialised_in_step_params(self) -> None:
        request = DocumentRequest(document_type="invoice", topic="Invoice", formats=[OutputFormat.PDF])
        workflow = build_document_workflow(request)
        params = workflow.steps[0].params["request"]
        self.assertEqual(params["document_type"], "invoice")
        self.assertEqual(params["formats"], ["pdf"])


# --------------------------------------------------------------------------- #
# 10. DocumentAgent
# --------------------------------------------------------------------------- #
class TestDocumentAgent(unittest.TestCase):
    def test_capabilities_registered(self) -> None:
        agent = DocumentAgent()
        agent.initialize({})
        capabilities = agent.list_capabilities()
        names = {c["name"] for c in capabilities}
        self.assertIn("generate_document", names)
        self.assertIn("list_document_types", names)
        self.assertIn("list_templates", names)

    def test_generate_capability(self) -> None:
        with TemporaryDirectory() as tmp:
            agent = DocumentAgent(generator=make_generator(Path(tmp)))
            agent.initialize({})
            response = agent.execute_capability(
                "generate_document",
                **{
                    "document_type": "invoice",
                    "topic": "Capability test invoice",
                    "formats": "docx",
                    "output_name": "capability_invoice",
                },
            )
            self.assertIn("capability_invoice", response["summary"])
            self.assertEqual(response["type"], "invoice")


if __name__ == "__main__":
    unittest.main()

