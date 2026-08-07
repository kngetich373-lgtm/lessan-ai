"""DOCX builder — the primary publishing-format exporter.

Renders a :class:`FormattedProject` with python-docx: title page, running
header, centred page-number footer, real Word TOC field, numbered headings,
auto-numbered figure/table captions, and publishing-standard body styles
(Times New Roman 12pt, 1.5 line spacing, 1" margins).
"""

from __future__ import annotations

import os
from typing import Any, List, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.logging import get_logger

from documents.builders.base import FormatBuilder
from documents.models import OutputFormat, Table

logger = get_logger("documents.docx")


def _set_run_font(
    run,
    name: str,
    size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> None:
    """Apply font attributes to a run, including east-asian fallback."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(str(color).lstrip("#"))


def _set_margins(section, margin_inches: float) -> None:
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Inches(margin_inches))


def _add_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    """Append a simple Word field (PAGE, TOC, …) to a paragraph."""
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = placeholder
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _restart_page_numbering(section) -> None:
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), "1")
    section._sectPr.append(pg_num_type)


def _heading_size(style, level: int) -> int:
    sizes = getattr(style, "heading_sizes", None) or [16, 14, 12, 12]
    return sizes[min(level - 1, len(sizes) - 1)]


class DocxBuilder(FormatBuilder):
    """Exports the document as a Word .docx file."""

    output_format = OutputFormat.DOCX

    def build(self, formatted, path: str) -> str:
        style = formatted.style
        document = Document()

        for section in document.sections:
            _set_margins(section, style.margin_inches)
        _configure_document_styles(document, style)

        # ------------------------------------------------------------------
        # Title page (first section) then body (new section, no header/footer
        # on the title page, page numbering restarts at 1).
        # ------------------------------------------------------------------
        if formatted.title_page:
            _render_title_page(document, formatted, style)
            body_section = document.add_section(WD_SECTION.NEW_PAGE)
            body_section.header.is_linked_to_previous = False
            body_section.footer.is_linked_to_previous = False
            _restart_page_numbering(body_section)
        else:
            body_section = document.sections[0]
            body_section.header.is_linked_to_previous = False
            body_section.footer.is_linked_to_previous = False
        _set_margins(body_section, style.margin_inches)

        if style.header_enabled:
            _add_header(body_section, formatted, style)
        _add_footer(body_section, style)

        # ------------------------------------------------------------------
        # Table of contents (real Word field)
        # ------------------------------------------------------------------
        if formatted.toc:
            heading = document.add_paragraph(style="Heading 1")
            _set_run_font(
                heading.add_run("Table of Contents"),
                style.font_family,
                size=_heading_size(style, 1),
                bold=style.heading_bold,
                color=style.heading_color,
            )
            toc_para = document.add_paragraph()
            _add_field(
                toc_para,
                r'TOC \o "1-3" \h \z \u',
                placeholder=(
                    "Table of contents — right-click this field and choose "
                    "'Update Field' in Microsoft Word."
                ),
            )

        for section in formatted.project.sections:
            self._render_section(document, section, formatted, style)

        document.save(path)
        return path

    # ------------------------------------------------------------------ #
    # Section rendering
    # ------------------------------------------------------------------ #
    def _render_section(self, document, section, formatted, style) -> None:
        number = section.meta.get("number", "")
        heading_text = f"{number}  {section.heading}".strip()
        heading = document.add_heading(level=min(4, section.level))
        _set_run_font(
            heading.add_run(heading_text),
            style.font_family,
            size=_heading_size(style, section.level),
            bold=style.heading_bold,
            color=style.heading_color,
        )
        heading.paragraph_format.space_before = Pt(18 if section.level == 1 else 12)
        heading.paragraph_format.space_after = Pt(8)
        heading.paragraph_format.line_spacing = 1.15

        for paragraph in section.paragraphs:
            para = document.add_paragraph()
            run = para.add_run(paragraph.text)
            _set_run_font(run, style.font_family, size=style.base_size)
            if paragraph.style == "quote":
                para.paragraph_format.left_indent = Inches(0.5)
                run.italic = True
            elif paragraph.style == "small":
                run.font.size = Pt(style.base_size - 1)

        for item in section.bullets:
            bullet = document.add_paragraph(style="List Bullet")
            run = bullet.add_run(item)
            _set_run_font(run, style.font_family, size=style.base_size)

        for index, item in enumerate(section.numbered, start=1):
            numbered = document.add_paragraph(style="List Number")
            run = numbered.add_run(item)
            _set_run_font(run, style.font_family, size=style.base_size)
            numbered.paragraph_format.left_indent = Inches(0.5)

        for table in section.tables:
            self._render_table(document, table, style)

        for figure in section.figures:
            self._render_figure(document, figure, style)

        for code_block in section.code_blocks:
            self._render_code(document, code_block, style)

        for reference in section.references:
            para = document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            run = para.add_run(reference)
            _set_run_font(run, style.font_family, size=style.base_size - 1)

        if section.page_break_after:
            document.add_page_break()

    def _render_table(self, document, table: Table, style) -> None:
        if table.caption:
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(table.caption)
            _set_run_font(
                run,
                style.font_family,
                size=style.base_size - 1,
                italic=style.caption_italic,
            )
            caption.paragraph_format.space_after = Pt(4)

        rows = table.rows
        if not rows:
            return
        column_count = max(len(row) for row in rows)
        doc_table = document.add_table(rows=len(rows), cols=column_count)
        doc_table.style = "Table Grid"
        doc_table.autofit = True

        for row_index, row in enumerate(rows):
            for col_index in range(column_count):
                cell = doc_table.cell(row_index, col_index)
                text = row[col_index] if col_index < len(row) else ""
                cell_para = cell.paragraphs[0]
                cell_para.text = ""
                run = cell_para.add_run(text)
                _set_run_font(
                    run,
                    style.font_family,
                    size=style.base_size - 1,
                    bold=(row_index == 0 and table.header_row),
                )

        spacing = document.add_paragraph()
        spacing.paragraph_format.space_after = Pt(style.space_after_pt)

    def _render_figure(self, document, figure, style) -> None:
        if figure.path and os.path.exists(str(figure.path)):
            try:
                document.add_picture(str(figure.path), width=Inches(5.5))
            except Exception as exc:  # noqa: BLE001
                logger.warning(f"Could not embed figure image {figure.path}: {exc}")
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(figure.caption)
        _set_run_font(
            run,
            style.font_family,
            size=style.base_size - 1,
            italic=style.caption_italic,
        )

    def _render_code(self, document, code_block, style) -> None:
        para = document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.line_spacing = 1.0
        lines = code_block.text.splitlines() or [""]
        for index, line in enumerate(lines):
            run = para.add_run(line)
            _set_run_font(run, style.mono_family, size=style.base_size - 1)
            if index < len(lines) - 1:
                run.add_break()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:fill"), "F2F2F2")
        para._p.get_or_add_pPr().append(shading)


def _render_title_page(document, formatted, style) -> None:
    """Centered title page: title, subtitle, author, date."""
    for _ in range(6):
        document.add_paragraph()

    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        title_para.add_run(formatted.title),
        style.font_family,
        size=28,
        bold=True,
        color=style.heading_color,
    )

    if formatted.subtitle:
        subtitle_para = document.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(
            subtitle_para.add_run(formatted.subtitle),
            style.font_family,
            size=16,
            italic=True,
        )

    for _ in range(8):
        document.add_paragraph()

    if formatted.author:
        author_para = document.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(author_para.add_run(formatted.author), style.font_family, size=14)

    date_para = document.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(date_para.add_run(formatted.date), style.font_family, size=12)


def _configure_document_styles(document, style) -> None:
    """Apply publishing defaults to Normal and Heading 1–4 styles."""
    normal = document.styles["Normal"]
    normal.font.name = style.font_family
    normal.font.size = Pt(style.base_size)
    normal.paragraph_format.line_spacing = style.line_spacing
    normal.paragraph_format.space_after = Pt(style.space_after_pt)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), style.font_family)

    for level in range(1, 5):
        heading_style = document.styles[f"Heading {level}"]
        heading_style.font.name = style.font_family
        heading_style.font.size = Pt(_heading_size(style, level))
        heading_style.font.bold = style.heading_bold
        heading_style.font.color.rgb = RGBColor.from_string(style.heading_color.lstrip("#"))
        heading_style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        heading_style.paragraph_format.space_after = Pt(8)
        heading_style.paragraph_format.line_spacing = 1.15
        heading_rpr = heading_style.element.get_or_add_rPr()
        heading_rfonts = heading_rpr.find(qn("w:rFonts"))
        if heading_rfonts is None:
            heading_rfonts = OxmlElement("w:rFonts")
            heading_rpr.append(heading_rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            heading_rfonts.set(qn(f"w:{attr}"), style.font_family)


def _add_header(section, formatted, style) -> None:
    """Running header: the document title, right-aligned, small caps feel."""
    header_text = style.header_text or formatted.title
    if not header_text:
        return
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(header_text)
    _set_run_font(run, style.font_family, size=10)
    _set_bottom_border(header)


def _add_footer(section, style) -> None:
    """Centred footer with the Lessan branding and a live page number."""
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if style.footer_text:
        run = footer.add_run(style.footer_text + "   ")
        _set_run_font(run, style.font_family, size=10)
    if style.page_numbers:
        _add_field(footer, "PAGE", placeholder="1")


def _set_bottom_border(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    ppr.append(borders)



